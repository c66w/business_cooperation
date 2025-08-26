"""
文档解析服务
"""

import os
import io
import tempfile
from typing import Dict, Any, Optional
from loguru import logger

# 文档解析库
import PyPDF2
from docx import Document
import openpyxl
from unstructured.partition.auto import partition

class DocumentParsingService:
    """文档解析服务类"""
    
    def __init__(self):
        """初始化解析服务"""
        self.supported_types = {
            'pdf': self._parse_pdf,
            'doc': self._parse_word,
            'docx': self._parse_word,
            'xls': self._parse_excel,
            'xlsx': self._parse_excel,
            'txt': self._parse_text
        }
        logger.info("文档解析服务初始化完成")
    
    async def parse_document(self, file_content: bytes, file_type: str, filename: str) -> Dict[str, Any]:
        """
        解析文档内容
        
        Args:
            file_content: 文件内容
            file_type: 文件类型
            filename: 文件名
            
        Returns:
            解析结果字典
        """
        try:
            # 检查文件类型是否支持
            if file_type.lower() not in self.supported_types:
                return {
                    "success": False,
                    "error": f"不支持的文件类型: {file_type}"
                }
            
            # 使用对应的解析方法
            parser = self.supported_types[file_type.lower()]
            content = await parser(file_content, filename)
            
            # 使用Unstructured进行结构化解析
            structured_content = await self._parse_with_unstructured(file_content, filename)
            
            return {
                "success": True,
                "raw_content": content,
                "structured_content": structured_content,
                "file_type": file_type,
                "content_length": len(content)
            }
            
        except Exception as e:
            logger.error(f"文档解析失败: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def _parse_pdf(self, file_content: bytes, filename: str) -> str:
        """解析PDF文件"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = []
            for page in pdf_reader.pages:
                text_content.append(page.extract_text())
            
            return "\n".join(text_content)
        except Exception as e:
            logger.error(f"PDF解析失败: {e}")
            raise
    
    async def _parse_word(self, file_content: bytes, filename: str) -> str:
        """解析Word文档"""
        try:
            # 创建临时文件
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                tmp_file.write(file_content)
                tmp_file.flush()
                
                # 解析文档
                doc = Document(tmp_file.name)
                text_content = []
                
                for paragraph in doc.paragraphs:
                    text_content.append(paragraph.text)
                
                # 解析表格
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            row_text.append(cell.text)
                        text_content.append(" | ".join(row_text))
                
                # 清理临时文件
                os.unlink(tmp_file.name)
                
                return "\n".join(text_content)
        except Exception as e:
            logger.error(f"Word文档解析失败: {e}")
            raise
    
    async def _parse_excel(self, file_content: bytes, filename: str) -> str:
        """解析Excel文件"""
        try:
            # 创建临时文件
            with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp_file:
                tmp_file.write(file_content)
                tmp_file.flush()
                
                # 解析工作簿
                workbook = openpyxl.load_workbook(tmp_file.name)
                text_content = []
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    text_content.append(f"工作表: {sheet_name}")
                    
                    for row in sheet.iter_rows(values_only=True):
                        row_text = []
                        for cell in row:
                            if cell is not None:
                                row_text.append(str(cell))
                        if row_text:
                            text_content.append(" | ".join(row_text))
                
                # 清理临时文件
                os.unlink(tmp_file.name)
                
                return "\n".join(text_content)
        except Exception as e:
            logger.error(f"Excel文件解析失败: {e}")
            raise
    
    async def _parse_text(self, file_content: bytes, filename: str) -> str:
        """解析文本文件"""
        try:
            # 尝试不同编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            
            for encoding in encodings:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # 如果所有编码都失败，使用错误处理
            return file_content.decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"文本文件解析失败: {e}")
            raise
    
    async def _parse_with_unstructured(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """使用Unstructured库进行结构化解析"""
        try:
            # 创建临时文件
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as tmp_file:
                tmp_file.write(file_content)
                tmp_file.flush()
                
                # 使用Unstructured解析
                elements = partition(filename=tmp_file.name)
                
                # 清理临时文件
                os.unlink(tmp_file.name)
                
                # 组织结构化内容
                structured = {
                    "titles": [],
                    "paragraphs": [],
                    "tables": [],
                    "lists": []
                }
                
                for element in elements:
                    element_type = str(type(element).__name__)
                    element_text = str(element)
                    
                    if "Title" in element_type:
                        structured["titles"].append(element_text)
                    elif "Table" in element_type:
                        structured["tables"].append(element_text)
                    elif "List" in element_type:
                        structured["lists"].append(element_text)
                    else:
                        structured["paragraphs"].append(element_text)
                
                return structured
                
        except Exception as e:
            logger.warning(f"Unstructured解析失败，使用基础解析: {e}")
            return {"paragraphs": [], "titles": [], "tables": [], "lists": []}

# 创建全局解析服务实例
parsing_service = DocumentParsingService()
